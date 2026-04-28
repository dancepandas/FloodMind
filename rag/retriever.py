"""
知识检索器

提供统一的知识检索接口，支持永久知识库和会话级知识库。
"""

import logging
import os
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import List, Optional, Dict, Any, Tuple

from langchain_core.documents import Document

from rag.embeddings import EmbeddingManager
from rag.vector_store import VectorStoreManager

logger = logging.getLogger(__name__)

SMALL_DOC_SIZE_THRESHOLD = 10000
SMALL_DOC_TOKEN_THRESHOLD = 3000


def _clean_filter_value(value: Any) -> Optional[str]:
    text = str(value or "").strip()
    return text or None


@dataclass
class DocumentProcessingResult:
    """文档处理结果"""
    success: bool
    method: str
    message: str
    doc_id: Optional[str] = None
    chunk_count: int = 0
    is_small_doc: bool = False
    content_preview: str = ""


@dataclass
class SearchResult:
    """检索结果"""
    documents: List[Document]
    scores: List[float]
    source: str
    query: str
    metadata_filter: Dict[str, Any] = field(default_factory=dict)

    @staticmethod
    def _format_display_source(doc: Document) -> str:
        display_path = str(doc.metadata.get("display_path", "") or "").strip()
        if display_path:
            return display_path

        relative_path = str(doc.metadata.get("relative_path", "") or "").strip()
        if relative_path:
            return f"./{relative_path}"

        filename = str(doc.metadata.get("filename", "") or "").strip()
        if filename:
            return filename

        return "未知来源"
    
    def to_context_text(self, max_length: int = 4000) -> str:
        """转换为上下文文本"""
        if not self.documents:
            return ""
        
        parts = []
        total_length = 0
        
        for i, doc in enumerate(self.documents):
            source = self._format_display_source(doc)
            meta_parts = []
            folder_path = str(doc.metadata.get("folder_path", "") or "").strip()
            asset_kind = str(doc.metadata.get("asset_kind", "") or "").strip()
            index_mode = str(doc.metadata.get("index_mode", "") or "").strip()
            if folder_path:
                meta_parts.append(f"目录: {folder_path}")
            if asset_kind:
                meta_parts.append(f"类型: {asset_kind}")
            if index_mode:
                meta_parts.append(f"索引: {index_mode}")
            meta_text = f" | {'; '.join(meta_parts)}" if meta_parts else ""
            chunk_text = f"【参考 {i+1}】(来源: {source}{meta_text})\n{doc.page_content}\n"
            
            if total_length + len(chunk_text) > max_length:
                break
            
            parts.append(chunk_text)
            total_length += len(chunk_text)
        
        return "\n".join(parts)


class KnowledgeRetriever:
    """知识检索器"""
    
    _instance: Optional['KnowledgeRetriever'] = None
    
    def __new__(cls, *args, **kwargs):
        if cls._instance is None:
            cls._instance = super().__new__(cls)
        return cls._instance
    
    def __init__(
        self,
        persist_dir: str = "./data/vector_store",
        embedding_model: str = "BAAI/bge-base-zh-v1.5",
        top_k: int = 5,
        small_doc_threshold: int = SMALL_DOC_SIZE_THRESHOLD,
    ):
        if hasattr(self, '_initialized') and self._initialized:
            return
            
        self.persist_dir = persist_dir
        self.top_k = top_k
        self.small_doc_threshold = small_doc_threshold
        
        self.embedding_manager = EmbeddingManager(model_name=embedding_model)
        
        self._permanent_store: Optional[VectorStoreManager] = None
        self._session_stores: Dict[str, VectorStoreManager] = {}
        
        self._small_docs: Dict[str, List[Document]] = {}
        
        self._initialized = True
        logger.info(f"知识检索器初始化: {persist_dir}, top_k={top_k}")
    
    @property
    def permanent_store(self) -> VectorStoreManager:
        """获取永久知识库"""
        if self._permanent_store is None:
            self._permanent_store = VectorStoreManager.get_permanent_store(
                self.persist_dir, self.embedding_manager
            )
        return self._permanent_store
    
    def get_session_store(self, session_id: str) -> VectorStoreManager:
        """获取会话级知识库"""
        if session_id not in self._session_stores:
            self._session_stores[session_id] = VectorStoreManager.get_session_store(
                self.persist_dir, session_id, self.embedding_manager
            )
        return self._session_stores[session_id]
    
    def is_small_document(self, content: str) -> bool:
        """判断是否为小文档"""
        if len(content) < self.small_doc_threshold:
            return True
        
        estimated_tokens = len(content) / 1.5
        if estimated_tokens < SMALL_DOC_TOKEN_THRESHOLD:
            return True
        
        return False
    
    def add_document(
        self,
        content: str,
        metadata: Optional[Dict[str, Any]] = None,
        session_id: Optional[str] = None,
        force_method: Optional[str] = None,
    ) -> DocumentProcessingResult:
        """
        添加文档到知识库
        
        Args:
            content: 文档内容
            metadata: 元数据
            session_id: 会话ID（可选，用于会话级存储）
            force_method: 强制指定处理方式 ("context" 或 "vector")
            
        Returns:
            DocumentProcessingResult: 处理结果
        """
        metadata = metadata or {}
        
        if "doc_id" not in metadata:
            import uuid
            metadata["doc_id"] = str(uuid.uuid4())
        metadata["created_at"] = datetime.now().isoformat()
        if session_id:
            metadata["session_id"] = session_id
        
        is_small = self.is_small_document(content)
        
        if force_method == "context" or (force_method is None and is_small):
            return self._add_to_context(content, metadata, session_id)
        elif force_method == "vector" or (force_method is None and not is_small):
            return self._add_to_vector_store(content, metadata, session_id)
        else:
            return self._add_to_context(content, metadata, session_id)
    
    def _add_to_context(
        self,
        content: str,
        metadata: Dict[str, Any],
        session_id: Optional[str],
    ) -> DocumentProcessingResult:
        """添加小文档到上下文"""
        doc = Document(page_content=content, metadata=metadata)
        
        key = session_id or "global"
        if key not in self._small_docs:
            self._small_docs[key] = []
        self._small_docs[key].append(doc)
        
        preview = content[:200] + "..." if len(content) > 200 else content
        
        logger.info(f"小文档已添加到上下文: {metadata.get('doc_id')}, 会话: {key}")
        
        return DocumentProcessingResult(
            success=True,
            method="context",
            message="文档已作为临时上下文添加（小文档）",
            doc_id=metadata.get("doc_id"),
            chunk_count=1,
            is_small_doc=True,
            content_preview=preview,
        )
    
    def _add_to_vector_store(
        self,
        content: str,
        metadata: Dict[str, Any],
        session_id: Optional[str],
    ) -> DocumentProcessingResult:
        """添加大文档到向量库"""
        try:
            if session_id:
                store = self.get_session_store(session_id)
            else:
                store = self.permanent_store
            
            ids = store.add_text(content, metadata)
            
            preview = content[:200] + "..." if len(content) > 200 else content
            
            logger.info(f"文档已添加到向量库: {metadata.get('doc_id')}, 分块数: {len(ids)}")
            
            return DocumentProcessingResult(
                success=True,
                method="vector",
                message=f"文档已添加到向量库，分为 {len(ids)} 个片段",
                doc_id=metadata.get("doc_id"),
                chunk_count=len(ids),
                is_small_doc=False,
                content_preview=preview,
            )
            
        except Exception as e:
            logger.error(f"添加文档到向量库失败: {e}")
            return DocumentProcessingResult(
                success=False,
                method="vector",
                message=f"添加失败: {str(e)}",
                doc_id=metadata.get("doc_id"),
            )
    
    def add_file(
        self,
        file_path: str,
        metadata: Optional[Dict[str, Any]] = None,
        session_id: Optional[str] = None,
    ) -> DocumentProcessingResult:
        """
        添加文件到知识库
        
        Args:
            file_path: 文件路径
            metadata: 元数据
            session_id: 会话ID
            
        Returns:
            DocumentProcessingResult: 处理结果
        """
        path = Path(file_path)
        if not path.exists():
            return DocumentProcessingResult(
                success=False,
                method="none",
                message=f"文件不存在: {file_path}",
            )
        
        metadata = metadata or {}
        metadata["source"] = str(path)
        metadata["filename"] = path.name
        metadata["file_type"] = path.suffix.lower()
        
        try:
            content = self._extract_file_content(path)
        except Exception as e:
            return DocumentProcessingResult(
                success=False,
                method="none",
                message=f"读取文件失败: {str(e)}",
            )
        
        return self.add_document(content, metadata, session_id)
    
    def _extract_file_content(self, path: Path) -> str:
        """提取文件内容"""
        suffix = path.suffix.lower()
        
        if suffix in [".txt", ".md", ".py", ".json", ".csv"]:
            return path.read_text(encoding="utf-8")
        
        if suffix == ".pdf":
            return self._extract_pdf(path)
        
        if suffix in [".docx", ".doc"]:
            return self._extract_docx(path)
        
        try:
            return path.read_text(encoding="utf-8")
        except Exception as e:
            raise ValueError(f"无法读取文件: {e}")
    
    def _extract_pdf(self, path: Path) -> str:
        """提取 PDF 内容"""
        try:
            import pypdf
            reader = pypdf.PdfReader(str(path))
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text
        except ImportError:
            raise ImportError("需要安装 pypdf: pip install pypdf")
    
    def _extract_docx(self, path: Path) -> str:
        """提取 Word 文档内容"""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(str(path))
            return "\n".join([para.text for para in doc.paragraphs])
        except ImportError:
            raise ImportError("需要安装 python-docx: pip install python-docx")
    
    def search(
        self,
        query: str,
        session_id: Optional[str] = None,
        top_k: Optional[int] = None,
        include_permanent: bool = True,
        include_session: bool = True,
        include_small_docs: bool = True,
        metadata_filter: Optional[Dict[str, Any]] = None,
    ) -> SearchResult:
        """
        检索相关知识
        
        Args:
            query: 查询文本
            session_id: 会话ID
            top_k: 返回结果数量
            include_permanent: 是否包含永久知识库
            include_session: 是否包含会话级知识库
            include_small_docs: 是否包含小文档上下文
            
        Returns:
            SearchResult: 检索结果
        """
        top_k = top_k or self.top_k
        merged_filter = self._merge_metadata_filter(query, metadata_filter)
        chromadb_filter = self._to_chromadb_where(merged_filter)
        all_docs: List[Tuple[Document, float, str]] = []
        
        if include_permanent:
            try:
                results = self.permanent_store.search_with_scores(query, k=top_k, filter=chromadb_filter)
                all_docs.extend((doc, score, "vector") for doc, score in results)
            except Exception as e:
                logger.warning(f"检索永久知识库失败: {e}")
        
        if include_session and session_id:
            try:
                session_store = self.get_session_store(session_id)
                results = session_store.search_with_scores(query, k=top_k, filter=chromadb_filter)
                all_docs.extend((doc, score, "vector") for doc, score in results)
            except Exception as e:
                logger.warning(f"检索会话知识库失败: {e}")
        
        if include_small_docs:
            small_docs = self._get_small_docs(session_id)
            for doc in small_docs:
                if not self._matches_metadata_filter(doc, merged_filter):
                    continue
                score = self._simple_similarity(query, doc.page_content)
                all_docs.append((doc, score, "small"))

        all_docs.sort(key=lambda x: self._to_relevance_score(x[1], x[2]), reverse=True)
        top_docs = all_docs[:top_k]

        documents = [doc for doc, _, _ in top_docs]
        scores = [self._to_relevance_score(score, score_type) for _, score, score_type in top_docs]
        
        source_parts = []
        if include_permanent:
            source_parts.append("永久知识库")
        if include_session and session_id:
            source_parts.append("会话知识库")
        if include_small_docs:
            source_parts.append("临时上下文")
        
        return SearchResult(
            documents=documents,
            scores=scores,
            source=", ".join(source_parts),
            query=query,
            metadata_filter=merged_filter,
        )
    
    def _get_small_docs(self, session_id: Optional[str]) -> List[Document]:
        """获取小文档列表"""
        docs = []
        
        if "global" in self._small_docs:
            docs.extend(self._small_docs["global"])
        
        if session_id and session_id in self._small_docs:
            docs.extend(self._small_docs[session_id])
        
        return docs

    @staticmethod
    def _to_relevance_score(raw_score: float, score_type: str) -> float:
        if score_type == "vector":
            try:
                score = max(float(raw_score), 0.0)
            except (TypeError, ValueError):
                return 0.0
            return 1.0 / (1.0 + score)
        try:
            return float(raw_score)
        except (TypeError, ValueError):
            return 0.0
    
    def _simple_similarity(self, query: str, content: str) -> float:
        """简单的文本相似度计算（用于小文档）"""
        query_terms = set(query.lower().split())
        content_terms = set(content.lower().split())
        
        if not query_terms or not content_terms:
            return 0.0
        
        intersection = query_terms & content_terms
        return len(intersection) / len(query_terms)

    def _normalize_filter(self, metadata_filter: Optional[Dict[str, Any]]) -> Dict[str, Any]:
        if not metadata_filter:
            return {}

        allowed_keys = {
            "asset_kind",
            "index_mode",
            "file_type",
            "folder_level_1",
            "folder_level_2",
            "folder_level_3",
            "folder_level_4",
            "folder_level_5",
            "folder_level_6",
            "filename",
        }
        normalized: Dict[str, Any] = {}
        for key, value in metadata_filter.items():
            if key not in allowed_keys:
                continue
            cleaned = _clean_filter_value(value)
            if cleaned is not None:
                normalized[key] = cleaned
        return normalized

    def _infer_filter_from_query(self, query: str) -> Dict[str, Any]:
        lower_query = str(query or "").lower()
        inferred: Dict[str, Any] = {}

        if any(token in lower_query for token in ["excel", "xlsx", "xls", "表格", "数据表", "sheet"]):
            inferred["asset_kind"] = "excel_asset"
            inferred["index_mode"] = "file_summary"
        elif any(token in lower_query for token in ["gis", "图层", "shp", "geojson", "地理信息", "栅格", "矢量"]):
            inferred["asset_kind"] = "gis_asset"
            inferred["index_mode"] = "file_summary"
        elif any(token in lower_query for token in ["图片", "图件", "照片", "示意图", "png", "jpg", "jpeg"]):
            inferred["asset_kind"] = "image_asset"
            inferred["index_mode"] = "file_summary"
        elif any(token in lower_query for token in ["word", "pdf", "报告", "文档", "说明", "方案", "正文"]):
            inferred["asset_kind"] = "text_document"
            inferred["index_mode"] = "content_chunk"

        return inferred

    def _merge_metadata_filter(self, query: str, metadata_filter: Optional[Dict[str, Any]]) -> Dict[str, Any]:
        explicit = self._normalize_filter(metadata_filter)
        inferred = self._infer_filter_from_query(query)
        return {**inferred, **explicit}

    @staticmethod
    def _to_chromadb_where(flat_filter: Dict[str, Any]) -> Optional[Dict[str, Any]]:
        if not flat_filter:
            return None
        if len(flat_filter) == 1:
            key, value = next(iter(flat_filter.items()))
            return {key: value}
        return {"$and": [{k: v} for k, v in flat_filter.items()]}

    @staticmethod
    def _matches_metadata_filter(doc: Document, metadata_filter: Dict[str, Any]) -> bool:
        if not metadata_filter:
            return True
        for key, expected in metadata_filter.items():
            actual = str(doc.metadata.get(key, "") or "").strip()
            if actual != str(expected):
                return False
        return True
    
    def get_context_for_query(
        self,
        query: str,
        session_id: Optional[str] = None,
        max_length: int = 4000,
        metadata_filter: Optional[Dict[str, Any]] = None,
    ) -> str:
        """
        获取查询相关的上下文文本
        
        Args:
            query: 查询文本
            session_id: 会话ID
            max_length: 最大长度
            
        Returns:
            格式化的上下文文本
        """
        result = self.search(query, session_id, metadata_filter=metadata_filter)
        return result.to_context_text(max_length)
    
    def clear_session(self, session_id: str):
        """清理会话数据"""
        if session_id in self._small_docs:
            del self._small_docs[session_id]
        
        VectorStoreManager.cleanup_session(self.persist_dir, session_id)
        
        if session_id in self._session_stores:
            del self._session_stores[session_id]
        
        logger.info(f"会话 {session_id} 的知识数据已清理")
    
    def get_small_docs_context(self, session_id: Optional[str] = None) -> str:
        """获取小文档的完整上下文（用于注入 Prompt）"""
        docs = self._get_small_docs(session_id)
        if not docs:
            return ""
        
        parts = []
        for i, doc in enumerate(docs):
            source = SearchResult._format_display_source(doc)
            parts.append(f"【文档 {i+1}】({source})\n{doc.page_content}")
        
        return "\n\n".join(parts)
    
    def get_stats(self, session_id: Optional[str] = None) -> Dict[str, Any]:
        """获取知识库统计信息"""
        stats = {
            "permanent_docs": self.permanent_store.get_document_count(),
            "small_docs_global": len(self._small_docs.get("global", [])),
        }
        
        if session_id:
            stats["session_docs"] = self.get_session_store(session_id).get_document_count()
            stats["small_docs_session"] = len(self._small_docs.get(session_id, []))
        
        return stats
    
    @classmethod
    def reset(cls):
        """重置单例（用于测试）"""
        cls._instance = None
