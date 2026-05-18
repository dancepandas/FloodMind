"""
向量数据库管理

使用 ChromaDB 作为向量存储，支持持久化和会话级隔离。
直接使用 chromadb 客户端，不依赖 LangChain。
"""

import logging
import os
import re
import shutil
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import chromadb

from agent.runtime.contracts.messages import Document
from rag.embeddings import EmbeddingManager
from rag.text_splitters import RecursiveCharacterTextSplitter, MarkdownHeaderTextSplitter

logger = logging.getLogger(__name__)

SMALL_DOC_THRESHOLD = 10000
SMALL_DOC_TOKEN_THRESHOLD = 3000


class VectorStoreManager:
    """向量数据库管理器"""

    _instances: Dict[str, 'VectorStoreManager'] = {}
    _reranker_cache: Dict[str, Any] = {}

    def __new__(
        cls,
        persist_dir: str,
        collection_name: str = "default",
        embedding_manager: Optional[EmbeddingManager] = None,
    ):
        cache_key = f"{persist_dir}:{collection_name}"
        if cache_key not in cls._instances:
            instance = super().__new__(cls)
            cls._instances[cache_key] = instance
        return cls._instances[cache_key]

    def __init__(
        self,
        persist_dir: str,
        collection_name: str = "default",
        embedding_manager: Optional[EmbeddingManager] = None,
    ):
        if hasattr(self, '_initialized') and self._initialized:
            return

        self.persist_dir = Path(persist_dir)
        self.collection_name = collection_name
        self.embedding_manager = embedding_manager or EmbeddingManager()

        self.persist_dir.mkdir(parents=True, exist_ok=True)

        self._client: Optional[chromadb.ClientAPI] = None
        self._collection = None
        self._initialized = True

        logger.info(f"向量存储管理器初始化: {persist_dir}, 集合: {collection_name}")

    @property
    def vectorstore(self):
        """获取 ChromaDB collection（懒加载，兼容旧属性名）"""
        if self._collection is None:
            self._init_chroma()
        return self._collection

    def _init_chroma(self):
        """初始化 ChromaDB 客户端和 collection"""
        self._client = chromadb.PersistentClient(path=str(self.persist_dir))
        self._collection = self._client.get_or_create_collection(
            name=self.collection_name,
            metadata={"hnsw:space": "cosine"},
        )
        logger.info(f"ChromaDB collection 就绪: {self.collection_name}, 文档数: {self._collection.count()}")

    def _embed_texts(self, texts: List[str]) -> List[List[float]]:
        """将文本列表转换为向量"""
        return self.embedding_manager.embed_documents(texts)

    def _embed_query(self, query: str) -> List[float]:
        """将查询文本转换为向量"""
        return self.embedding_manager.embed_query(query)

    def add_documents(
        self,
        documents: List[Document],
        metadatas: Optional[List[Dict[str, Any]]] = None,
    ) -> List[str]:
        """添加文档到向量库"""
        if not documents:
            return []

        if metadatas:
            for i, meta in enumerate(metadatas):
                if i < len(documents):
                    documents[i].metadata.update(meta)

        for doc in documents:
            if "doc_id" not in doc.metadata:
                doc.metadata["doc_id"] = str(uuid.uuid4())
            if "created_at" not in doc.metadata:
                doc.metadata["created_at"] = datetime.now().isoformat()

        collection = self.vectorstore
        ids = [doc.metadata.get("doc_id", str(uuid.uuid4())) for doc in documents]
        texts = [doc.page_content for doc in documents]
        embeddings = self._embed_texts(texts)
        metas = [doc.metadata for doc in documents]

        # ChromaDB metadata values must be str, int, float, or bool
        clean_metas = []
        for meta in metas:
            clean = {}
            for k, v in meta.items():
                if isinstance(v, (str, int, float, bool)):
                    clean[k] = v
                else:
                    clean[k] = str(v)
            clean_metas.append(clean)

        collection.add(
            ids=ids,
            documents=texts,
            embeddings=embeddings,
            metadatas=clean_metas,
        )

        logger.info(f"添加 {len(documents)} 个文档到向量库，集合: {self.collection_name}")
        return ids

    def add_text(
        self,
        text: str,
        metadata: Optional[Dict[str, Any]] = None,
        chunk_size: int = 300,
        chunk_overlap: int = 100,
    ) -> List[str]:
        """添加文本到向量库（自动分块，表格感知）"""
        if not text.strip():
            return []

        metadata = metadata or {}

        raw_chunks = self._split_preserving_tables(text, chunk_size, chunk_overlap)

        documents = [
            Document(page_content=chunk, metadata={**metadata, "chunk_index": i})
            for i, chunk in enumerate(raw_chunks)
        ]

        return self.add_documents(documents)

    def _split_preserving_tables(
        self,
        text: str,
        chunk_size: int,
        chunk_overlap: int,
    ) -> List[str]:
        """表格感知分块核心方法"""
        recursive_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", "。", "！", "？", "；", "，", " ", ""],
        )

        lines = text.split("\n")
        result_chunks: List[str] = []

        non_table_buf: List[str] = []
        table_buf: List[str] = []

        def is_table_title(line: str) -> bool:
            stripped = line.strip()
            if not stripped:
                return False
            if re.match(r"^#{1,6}\s+", stripped):
                return True
            title_keywords = ["表", "Table", "TABLE", "表格", "列表", "清单"]
            if any(kw in stripped for kw in title_keywords):
                if len(stripped) < 100:
                    return True
            if len(stripped) < 50 and not stripped.endswith(("。", "！", "？", ".", "!", "?")):
                return True
            return False

        def extract_table_titles() -> List[str]:
            titles = []
            i = len(non_table_buf) - 1
            while i >= 0:
                line = non_table_buf[i]
                if not line.strip():
                    titles.insert(0, line)
                    i -= 1
                elif is_table_title(line):
                    titles.insert(0, line)
                    i -= 1
                else:
                    break
            return titles

        def flush_non_table() -> None:
            if non_table_buf:
                segment = "\n".join(non_table_buf).strip()
                if segment:
                    result_chunks.extend(recursive_splitter.split_text(segment))
                non_table_buf.clear()

        def flush_table() -> None:
            if table_buf:
                table_text = "\n".join(table_buf).strip()
                if table_text:
                    result_chunks.append(table_text)
                table_buf.clear()

        for line in lines:
            if re.match(r"^\s*\|", line):
                if not table_buf:
                    titles = extract_table_titles()
                    if titles:
                        for title in titles:
                            non_table_buf.remove(title)
                        table_buf.extend(titles)
                    flush_non_table()
                table_buf.append(line)
            else:
                if table_buf:
                    flush_table()
                non_table_buf.append(line)

        flush_non_table()
        flush_table()

        return result_chunks

    def add_file(
        self,
        file_path: str,
        metadata: Optional[Dict[str, Any]] = None,
        chunk_size: int = 300,
        chunk_overlap: int = 100,
    ) -> List[str]:
        """添加文件到向量库"""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        metadata = metadata or {}
        metadata["source"] = str(path)
        metadata["filename"] = path.name
        metadata["file_type"] = path.suffix.lower()

        if path.suffix.lower() == ".md":
            return self._add_markdown_file(path, metadata, chunk_size, chunk_overlap)

        text = self._extract_text(path)
        return self.add_text(text, metadata, chunk_size, chunk_overlap)

    def _add_markdown_file(
        self,
        path: Path,
        metadata: Dict[str, Any],
        chunk_size: int,
        chunk_overlap: int,
    ) -> List[str]:
        """Markdown 文件两阶段 + 表格感知分块"""
        text = path.read_text(encoding="utf-8")
        if not text.strip():
            return []

        header_splitter = MarkdownHeaderTextSplitter(
            headers_to_split_on=[
                ("#", "h1"),
                ("##", "h2"),
                ("###", "h3"),
            ],
            strip_headers=False,
        )
        header_docs = header_splitter.split_text(text)

        final_docs = []
        chunk_index = 0
        for doc in header_docs:
            if len(doc.page_content) <= chunk_size:
                doc.metadata.update({**metadata, "chunk_index": chunk_index})
                final_docs.append(doc)
                chunk_index += 1
            else:
                sub_chunks = self._split_preserving_tables(
                    doc.page_content, chunk_size, chunk_overlap
                )
                for sub in sub_chunks:
                    final_docs.append(
                        Document(
                            page_content=sub,
                            metadata={**metadata, **doc.metadata, "chunk_index": chunk_index},
                        )
                    )
                    chunk_index += 1

        logger.info(
            f"Markdown 文件 {path.name} 分块完成："
            f"标题节数={len(header_docs)}，最终块数={len(final_docs)}"
        )
        return self.add_documents(final_docs)

    def _extract_text(self, path: Path) -> str:
        """从文件中提取文本"""
        suffix = path.suffix.lower()

        if suffix in [".txt", ".md", ".py", ".json", ".csv"]:
            return path.read_text(encoding="utf-8")

        if suffix == ".pdf":
            return self._extract_pdf(path)

        if suffix in [".docx", ".doc"]:
            return self._extract_docx(path)

        try:
            return path.read_text(encoding="utf-8")
        except Exception as e:
            logger.warning(f"无法读取文件 {path}: {e}")
            return ""

    def _extract_pdf(self, path: Path) -> str:
        try:
            import pypdf
            reader = pypdf.PdfReader(str(path))
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text
        except ImportError:
            logger.warning("未安装 pypdf，无法读取 PDF")
            return ""
        except Exception as e:
            logger.error(f"读取 PDF 失败: {e}")
            return ""

    def _extract_docx(self, path: Path) -> str:
        suffix = path.suffix.lower()
        if suffix == ".doc":
            return self._extract_doc(path)

        try:
            from docx import Document as DocxDocument
            from docx.table import Table
            from docx.text.paragraph import Paragraph
            from docx.oxml.ns import qn

            doc = DocxDocument(str(path))
            parts = []

            def iter_block_items(parent):
                if hasattr(parent, 'element'):
                    parent_elm = parent.element.body
                else:
                    parent_elm = parent
                for child in parent_elm.iterchildren():
                    if child.tag == qn('w:p'):
                        yield Paragraph(child, parent)
                    elif child.tag == qn('w:tbl'):
                        yield Table(child, parent)

            def table_to_markdown(table: Table) -> str:
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                    rows.append(cells)
                if not rows:
                    return ""
                col_count = len(rows[0])
                md_lines = []
                header = "| " + " | ".join(rows[0]) + " |"
                md_lines.append(header)
                separator = "| " + " | ".join(["---"] * col_count) + " |"
                md_lines.append(separator)
                for row in rows[1:]:
                    while len(row) < col_count:
                        row.append("")
                    md_lines.append("| " + " | ".join(row[:col_count]) + " |")
                return "\n".join(md_lines)

            for block in iter_block_items(doc):
                if isinstance(block, Paragraph):
                    text = block.text.strip()
                    if text:
                        parts.append(text)
                elif isinstance(block, Table):
                    table_md = table_to_markdown(block)
                    if table_md:
                        parts.append(table_md)

            return "\n\n".join(parts)
        except Exception as e:
            logger.error(f"读取 Word 文档失败: {e}")
            return ""

    def _extract_doc(self, path: Path) -> str:
        try:
            import platform
            if platform.system() != 'Windows':
                return self._extract_doc_via_textract(path)
            try:
                import win32com.client
                import pythoncom
                pythoncom.CoInitialize()
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                doc = word.Documents.Open(str(path.absolute()))
                text = doc.Content.Text
                doc.Close(False)
                word.Quit()
                pythoncom.CoUninitialize()
                return text
            except ImportError:
                return self._extract_doc_via_textract(path)
            except Exception as e:
                logger.error(f"使用 Word COM 读取 .doc 失败: {e}")
                return self._extract_doc_via_textract(path)
        except Exception as e:
            logger.error(f"读取 .doc 文档失败: {e}")
            return ""

    def _extract_doc_via_textract(self, path: Path) -> str:
        try:
            import textract
            text = textract.process(str(path))
            return text.decode('utf-8') if isinstance(text, bytes) else text
        except ImportError:
            logger.warning("未安装 textract，无法读取 .doc 文件")
            return ""
        except Exception as e:
            logger.error(f"使用 textract 读取 .doc 失败: {e}")
            return ""

    def search(
        self,
        query: str,
        k: int = 5,
        filter: Optional[Dict[str, Any]] = None,
    ) -> List[Document]:
        """相似度检索"""
        collection = self.vectorstore
        if collection.count() == 0:
            return []

        query_embedding = self._embed_query(query)
        query_params: Dict[str, Any] = {
            "query_embeddings": [query_embedding],
            "n_results": min(k, collection.count()),
        }
        if filter:
            query_params["where"] = filter

        results = collection.query(**query_params)

        documents = []
        if results and results.get("documents"):
            for i, text in enumerate(results["documents"][0]):
                meta = {}
                if results.get("metadatas") and results["metadatas"][0]:
                    meta = results["metadatas"][0][i] or {}
                documents.append(Document(page_content=text, metadata=meta))

        logger.info(f"检索查询: {query[:50]}..., 返回 {len(documents)} 个结果")
        return documents

    def search_with_scores(
        self,
        query: str,
        k: int = 5,
        filter: Optional[Dict[str, Any]] = None,
    ) -> List[tuple]:
        """带分数的相似度检索"""
        collection = self.vectorstore
        if collection.count() == 0:
            return []

        query_embedding = self._embed_query(query)
        query_params: Dict[str, Any] = {
            "query_embeddings": [query_embedding],
            "n_results": min(k, collection.count()),
        }
        if filter:
            query_params["where"] = filter

        results = collection.query(**query_params)

        documents_with_scores = []
        if results and results.get("documents"):
            for i, text in enumerate(results["documents"][0]):
                meta = {}
                if results.get("metadatas") and results["metadatas"][0]:
                    meta = results["metadatas"][0][i] or {}
                score = 0.0
                if results.get("distances") and results["distances"][0]:
                    score = results["distances"][0][i]
                documents_with_scores.append((
                    Document(page_content=text, metadata=meta),
                    score,
                ))

        return documents_with_scores

    def search_with_rerank(
        self,
        query: str,
        k: int = 5,
        fetch_k: int = 50,
        filter: Optional[Dict[str, Any]] = None,
        reranker_model: str = "BAAI/bge-reranker-base",
    ) -> List[Document]:
        """两阶段检索：向量召回 → Cross-Encoder 重排序"""
        fetch_k = max(fetch_k, k)
        candidates = self.search(query, k=fetch_k, filter=filter)

        if not candidates:
            return []

        logger.info(
            f"两阶段检索 — 阶段一召回: {len(candidates)} 条候选，"
            f"查询: {query[:50]}..."
        )

        reranker = self._get_reranker(reranker_model)
        pairs = [(query, doc.page_content) for doc in candidates]
        scores = reranker.predict(pairs)

        ranked = sorted(
            zip(scores, candidates),
            key=lambda x: x[0],
            reverse=True,
        )
        top_docs = [doc for _, doc in ranked[:k]]

        logger.info(
            f"两阶段检索 — 阶段二精排完成，返回 top-{len(top_docs)} 条"
        )
        return top_docs

    def _get_reranker(self, model_name: str) -> Any:
        """获取 Cross-Encoder 重排序器（懒加载 + 类级缓存）"""
        if model_name not in VectorStoreManager._reranker_cache:
            try:
                from sentence_transformers import CrossEncoder
            except ImportError as exc:
                raise ImportError(
                    "重排序功能需要 sentence-transformers，"
                    "请执行: pip install sentence-transformers"
                ) from exc

            logger.info(f"加载 Cross-Encoder 重排序模型: {model_name}")
            VectorStoreManager._reranker_cache[model_name] = CrossEncoder(
                model_name,
                max_length=512,
            )
            logger.info(f"Cross-Encoder 模型加载完成: {model_name}")

        return VectorStoreManager._reranker_cache[model_name]

    def delete_by_metadata(self, filter: Dict[str, Any]) -> int:
        """按元数据删除文档"""
        try:
            collection = self.vectorstore
            if filter:
                where_clause = self._build_where_clause(filter)
                result = collection.delete(where=where_clause)
                deleted_count = len(result.get("ids", []))
            else:
                result = collection.delete()
                deleted_count = len(result.get("ids", []))

            logger.info(f"删除 {deleted_count} 个文档，过滤条件: {filter}")
            return deleted_count
        except Exception as e:
            logger.error(f"删除文档失败: {e}")
            return 0

    def _build_where_clause(self, filter: Dict[str, Any]) -> Dict[str, Any]:
        """构建 ChromaDB where 子句"""
        if len(filter) == 1:
            key, value = next(iter(filter.items()))
            return {key: value}
        elif len(filter) > 1:
            return {"$and": [{k: v} for k, v in filter.items()]}
        return {}

    def clear(self):
        """清空当前集合"""
        self.delete_by_metadata({})
        logger.info(f"集合 {self.collection_name} 已清空")

    def delete_collection(self):
        """删除整个集合"""
        try:
            if self._client and self._collection:
                self._client.delete_collection(self.collection_name)
                self._collection = None
                logger.info(f"集合 {self.collection_name} 已删除")
        except Exception as e:
            logger.error(f"删除集合失败: {e}")

    def get_document_count(self) -> int:
        """获取文档数量"""
        try:
            return self.vectorstore.count()
        except Exception:
            return 0

    @classmethod
    def get_session_store(
        cls,
        base_dir: str,
        session_id: str,
        embedding_manager: Optional[EmbeddingManager] = None,
    ) -> 'VectorStoreManager':
        """获取会话级向量存储"""
        persist_dir = os.path.join(base_dir, "sessions", session_id)
        return cls(
            persist_dir=persist_dir,
            collection_name=f"session_{session_id}",
            embedding_manager=embedding_manager,
        )

    @classmethod
    def get_permanent_store(
        cls,
        base_dir: str,
        embedding_manager: Optional[EmbeddingManager] = None,
    ) -> 'VectorStoreManager':
        """获取永久向量存储"""
        persist_dir = os.path.join(base_dir, "permanent")
        return cls(
            persist_dir=persist_dir,
            collection_name="permanent",
            embedding_manager=embedding_manager,
        )

    @classmethod
    def cleanup_session(cls, base_dir: str, session_id: str):
        """清理会话级向量存储"""
        session_dir = Path(base_dir) / "sessions" / session_id
        if session_dir.exists():
            shutil.rmtree(session_dir)
            logger.info(f"已清理会话向量存储: {session_id}")

    @classmethod
    def reset_all(cls):
        """重置所有实例（用于测试）"""
        cls._instances.clear()